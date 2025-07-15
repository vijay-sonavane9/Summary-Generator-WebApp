import re
import fitz  # PyMuPDF
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import WordNetLemmatizer
from transformers import BartTokenizer, BartForConditionalGeneration
from docx import Document
from fpdf import FPDF
import torch

# Download NLTK resources
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using PyMuPDF (fitz)"""
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text

def preprocess_text(text):
    """Clean and preprocess text using NLTK"""
    # Basic cleaning
    text = text.lower()
    text = re.sub(r'[^a-zA-Z0-9\s.,;:!?]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Advanced NLP preprocessing
    sentences = sent_tokenize(text)
    processed_sentences = []
    
    lemmatizer = WordNetLemmatizer()
    stop_words = set(stopwords.words('english'))
    
    for sent in sentences:
        tokens = word_tokenize(sent)
        filtered_tokens = [
            lemmatizer.lemmatize(token)
            for token in tokens
            if token not in stop_words and len(token) > 2
        ]
        processed_sentences.append(" ".join(filtered_tokens))
    
    return " ".join(processed_sentences)

def summarize_text(text, model, tokenizer, device):
    """Generate summary using BART model"""
    # Tokenize input
    inputs = tokenizer(
        [text],
        max_length=1024,
        truncation=True,
        return_tensors="pt"
    ).to(device)
    
    # Generate summary
    summary_ids = model.generate(
        inputs.input_ids,
        num_beams=4,
        max_length=150,
        min_length=40,
        early_stopping=True,
        no_repeat_ngram_size=3
    )
    
    # Decode output
    summary = tokenizer.decode(
        summary_ids[0],
        skip_special_tokens=True,
        clean_up_tokenization_spaces=True
    )
    return summary

def save_summary(summary, output_path, format_type):
    """Save summary in specified format"""
    format_type = format_type.lower()
    
    if format_type == 'txt':
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(summary)
            
    elif format_type == 'docx':
        doc = Document()
        doc.add_heading('Document Summary', 0)
        doc.add_paragraph(summary)
        doc.save(output_path)
        
    elif format_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Document Summary", ln=True, align='C')
        pdf.ln(10)
        
        # Add summary content
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 8, txt=summary)
        pdf.output(output_path)
        
    else:
        raise ValueError("Unsupported format. Use txt, docx, or pdf")

def main():
    # User inputs
    pdf_path = input("Enter PDF file path: ").strip()
    output_path = input("Enter output file path: ").strip()
    format_type = input("Output format (txt/docx/pdf): ").strip().lower()

    # Device configuration
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"Using device: {device}")
    
    # Load BART model and tokenizer
    print("Loading BART model...")
    tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')
    model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn').to(device)
    
    # Process document
    print("Extracting text from PDF...")
    raw_text = extract_text_from_pdf(pdf_path)
    
    print("Preprocessing text...")
    cleaned_text = preprocess_text(raw_text)
    
    # Handle empty text
    if not cleaned_text.strip():
        raise ValueError("No text extracted from PDF")
    
    print("Generating summary...")
    summary = summarize_text(cleaned_text, model, tokenizer, device)
    
    print(f"Saving summary as {format_type.upper()}...")
    save_summary(summary, output_path, format_type)
    
    print(f"Summary saved to {output_path}")
    print(f"Summary length: {len(summary)} characters")
    print("\nSummary Preview:")
    print(summary[:300] + "..." if len(summary) > 300 else summary)

if _name_ == "_main_":
    main()
